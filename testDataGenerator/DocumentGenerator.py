from faker import Faker
from docx.shared import Pt, Cm
from docx.shared import Inches
import random


class DocumentGenerator:
    def __init__(self):
        self.document = Document()
        self.faker = Faker('ru_RU')

    def set_font_size(self, size):
        style = self.document.styles['Normal']
        font = style.font
        font.size = Pt(size)

    # Генерация углового бланка по ГОСТ
    def generate_corner_stamp(self):
        company_name = self.faker.company()
        address = self.faker.address()
        phone = self.faker.phone_number()

        # ИНН - 10цифр, ОГРН -13 цифр, ОКПО - 8 цифр
        inn = self.faker.unique.random_number(digits=10)
        ogrn = self.faker.unique.random_number(digits=13)
        okpo = self.faker.unique.random_number(digits=8)

        document_number = "№"
        document_date = "_"

        # Путь к изображению герба
        image_path = 'test.jpg'
        self.document.add_picture(image_path, width=Inches(random.uniform(0.5, 1)))

        paragraph = self.document.add_paragraph()
        paragraph.width = Cm(random.randint(5, 8))
        paragraph.add_run(company_name).bold = True
        paragraph.add_run('\n'*random.randint(1, 3) + address).italic = True
        paragraph.add_run('\n' + phone).italic = True
        paragraph.add_run('\n'*random.randint(1, 2) + 'ИНН: ' + str(inn)).italic = True
        paragraph.add_run('\n'*random.randint(0, 1) + 'ОГРН: ' + str(ogrn)).italic = True
        paragraph.add_run('\nОКПО: ' + str(okpo)).italic = True
        paragraph.add_run('\n'*random.randint(1, 2) + document_date * random.randint(3, 8)).italic = True
        paragraph.add_run(document_number + '_' * random.randint(3, 8)).italic = True

    # генерация страниц с угловым бланком
    def generate_pages(self, num_pages):
        for _ in range(num_pages):
            self.set_font_size(random.randint(12, 14))
            self.generate_corner_stamp()

            self.document.add_paragraph('\n\n\n')
            title = self.document.add_heading('Заголовок документа', level=1)
            title.alignment = 1

            text = self.faker.paragraphs(random.randint(2, 3))
            self.document.add_paragraph('\n' + text)
            text = 'Документ был создан с использованием программных средств.'
            self.document.add_paragraph('\n' + text)

            self.document.add_paragraph('\n\n\n')
            author = self.faker.name()
            fio_paragraph = self.document.add_paragraph(author)
            fio_paragraph.alignment = 2
            self.document.add_page_break()

    def save_document(self, filename):
        self.document.save(filename)

